from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Bharath_Reddy_Kalva_Uber_Conversational_Assistants_Resume.docx"


def set_margins(section):
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def shade_heading(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF1FB")
    p_pr.append(shd)


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F5597")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 78, 121)
    add_rule(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8.65)
    return p


def add_role(doc, left, right):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.6)
    table.columns[1].width = Inches(1.85)
    remove_table_borders(table)
    set_cell_text(table.cell(0, 0), left, bold=True)
    set_cell_text(table.cell(0, 1), right, bold=True)
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(4.6 if cell == row.cells[0] else 1.85)


def add_project(doc, title, tech, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r2 = p.add_run(" | " + tech)
    r2.font.name = "Calibri"
    r2.font.size = Pt(8.7)
    r2.font.color.rgb = RGBColor(89, 89, 89)
    for b in bullets:
        add_bullet(doc, b)


def add_hyperlink(paragraph, text, url, size=8.7):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    font.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(font)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_project_with_link(doc, title, url, tech, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    sep = p.add_run(" | ")
    sep.font.name = "Calibri"
    sep.font.size = Pt(8.7)
    add_hyperlink(p, url.replace("https://", "").rstrip("/"), url)
    r2 = p.add_run(" | " + tech)
    r2.font.name = "Calibri"
    r2.font.size = Pt(8.7)
    r2.font.color.rgb = RGBColor(89, 89, 89)
    for b in bullets:
        add_bullet(doc, b)


doc = Document()
set_margins(doc.sections[0])

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9)
styles["Normal"].paragraph_format.space_after = Pt(1)
styles["Normal"].paragraph_format.line_spacing = 1.0

for style_name in ["List Bullet"]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(8.65)
    style.paragraph_format.space_after = Pt(1.2)
    style.paragraph_format.line_spacing = 1.0

name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(0)
run = name.add_run("BHARATH REDDY KALVA")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(31, 78, 121)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(3)
cr = contact.add_run(
    "+1 (716) 400-3422 | bharathreddykalva10@gmail.com | linkedin.com/in/bharathreddy-kalva | github.com/kalva | Buffalo, NY"
)
cr.font.name = "Calibri"
cr.font.size = Pt(8.6)

add_section(doc, "PROFESSIONAL SUMMARY")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
summary = (
    "Software Engineer with 2+ years of experience building scalable Java/Spring Boot microservices, "
    "customer-facing platforms, and GenAI-powered chatbot and assistant workflows. Shipped production services for a "
    "communication platform supporting 300,000+ daily agent logins across 150 countries with 99.9% availability. "
    "Experienced in REST APIs, scalable system architecture, distributed systems, LLM prompt integration, conversational "
    "flows, observability, CI/CD, AWS deployments, testing, and cross-functional delivery with product, data science, QA, "
    "and design teams."
)
sr = p.add_run(summary)
sr.font.name = "Calibri"
sr.font.size = Pt(8.9)

add_section(doc, "TECHNICAL SKILLS")
skills = [
    ("Languages", "Java, Python, JavaScript, TypeScript, SQL, C++, C, Go (learning)"),
    ("Backend", "Spring Boot, REST APIs, Microservices, Node.js, NestJS, FastAPI, OAuth2, JWT, RBAC, scalable services"),
    ("AI / GenAI", "LLM integration, prompt engineering, conversational AI, chatbot workflows, assistant platforms, OpenAI, Groq, NLU/NLP features"),
    ("Data / Messaging", "PostgreSQL, MySQL, Redis, Prisma ORM, indexing, query optimization, BullMQ, WebSockets"),
    ("Cloud / DevOps", "AWS EC2, RDS, S3, CloudWatch, Docker, GitHub Actions, Jenkins, Vercel, Render, Linux"),
    ("Quality", "JUnit, Mockito, Jest, unit/integration testing, code reviews, logging, monitoring, performance troubleshooting, Agile/Scrum, technical documentation"),
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0.5)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(8.65)
    v = p.add_run(value)
    v.font.name = "Calibri"
    v.font.size = Pt(8.65)

add_section(doc, "PROFESSIONAL EXPERIENCE")
add_role(doc, "Software Engineer - Ozonetel (oneCXI)", "Aug 2022 - Jun 2024")
experience_bullets = [
    "Architected and delivered Java/Spring Boot REST microservices for a cloud communication and customer-support platform used by 300,000+ daily agent logins across 150 countries with 99.9% availability.",
    "Built extensible backend services for agent/customer workflows using OOP design patterns, clean API contracts, PostgreSQL persistence, and service-layer validation to support high-volume operational use cases.",
    "Reduced API response latency by 35% through SQL query tuning, composite indexing, and backend refactoring, improving real-time user workflows for support and operations teams.",
    "Implemented JWT authentication, RBAC authorization, and Spring Security controls; added JUnit/Mockito unit and integration tests, reducing security defects by 40% and improving code-review confidence.",
    "Deployed and monitored microservices on AWS EC2 with CloudWatch alerts, logs, and operational dashboards, reducing unplanned downtime by 25% through faster issue detection and troubleshooting.",
    "Collaborated with product managers, data scientists, QA, and design teams across 10+ Agile sprint cycles to translate customer requirements into maintainable platform features and release documentation.",
]
for b in experience_bullets:
    add_bullet(doc, b)

add_role(doc, "Software Engineer Intern - Ozonetel", "Jun 2022 - Aug 2022")
intern_bullets = [
    "Designed and deployed Jenkins, Git, and Docker CI/CD pipelines across development, staging, and production, reducing deployment time by 75% and eliminating manual release errors.",
    "Documented the release workflow and onboarded 2 engineers, reducing average release preparation time from 4 hours to under 30 minutes.",
]
for b in intern_bullets:
    add_bullet(doc, b)

add_section(doc, "TECHNICAL PROJECTS")
add_project_with_link(
    doc,
    "PackMind - Conversational AI Travel Packing Assistant",
    "https://packmind-two.vercel.app/",
    "Next.js, TypeScript, Supabase, Clerk, Groq AI, Twilio, OpenWeather, Vercel",
    [
        "Built a deployed GenAI assistant that guides users through trip creation and generates personalized, weather-aware packing plans using LLM prompts, destination context, trip type, activities, and travel style.",
        "Engineered a conversational flow with authenticated multi-user persistence, AI destination intelligence, and WhatsApp/SMS notifications via Twilio for scheduled packing reminders.",
        "Integrated OpenWeather and LLM-generated local guidance to produce context-aware recommendations, cultural notes, and resource warnings for end users.",
    ],
)
add_project(
    doc,
    "Distributed Multi-Client Chat System",
    "JavaScript/Node.js, WebSockets, REST APIs, Distributed Systems",
    [
        "Built a real-time multi-client chat application with concurrent client sessions, message routing, connection lifecycle handling, and resilient user-to-user communication flows.",
        "Designed the system around scalable service boundaries and API contracts relevant to customer-facing conversational platforms and support-channel interfaces.",
    ],
)
add_project_with_link(
    doc,
    "Research Brain - Proactive AI Research Assistant",
    "https://research-brain.butterbase.dev/login",
    "Butterbase, RocketRide, AI, Automation, Slack, Discord",
    [
        "Created an AI research assistant that tracks topics, remembers previously learned context, and proactively sends updates through Slack and Discord channels.",
        "Implemented assistant-style memory and notification workflows to convert ongoing research signals into timely, human-readable updates for users.",
    ],
)
add_project(
    doc,
    "HireSense AI - Resume and Job Description Intelligence Platform",
    "React, FastAPI, OpenAI, Python, Vercel, Render",
    [
        "Built a full-stack LLM application that parses PDF/DOCX/TXT/Markdown resumes, compares them against job descriptions, and returns match scores, skill gaps, ATS feedback, rewrites, and interview preparation.",
        "Designed a production-ready analyzer with OpenAI-powered responses and deterministic fallback behavior so the application remains usable when AI provider calls fail.",
    ],
)

add_section(doc, "EDUCATION")
add_role(doc, "University at Buffalo (SUNY) - M.S. in Computer Science", "Aug 2024 - Dec 2025")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Relevant Coursework: Distributed Systems, Operating Systems, Database Systems, Algorithms, Computer Networks")
r.font.name = "Calibri"
r.font.size = Pt(8.65)
add_role(doc, "S.N.I.S.T - Bachelor of Engineering in Electrical and Computer Engineering", "Jul 2022")

doc.save(OUT)
print(OUT)
